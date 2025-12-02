import docx
import docx.enum.text
from docx.shared import Cm, Mm, Pt
from docx.enum.text import WD_LINE_SPACING

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

section = doc.sections[-1]
section.top_margin = Cm(2) #Верхний отступ
section.bottom_margin = Cm(2)#Нижний отступ
section.left_margin = Cm(3) #Отступ слева
section.right_margin = Cm(1.5) #Отступ справа

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run("Полное название организации")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n\n\n\n\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
add = text.add_run("Название кафедры")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n\n\n\n\n\n\n\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run("Отчет по лабораторной работе №#")
add.font.size = Pt(24)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run("по дисциплине «Название дисциплины»")
add.font.size = Pt(16)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n")
add.font.size = Pt(16)
add = text.add_run("\n\n")
add.font.size = Pt(14)
add = text.add_run("\n\n\n\n")
add.font.size = Pt(18)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("                                                           ")
add.font.size = Pt(18)
add = text.add_run("Выполнил: студент группы Номер ")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("                                                                                                Фамилия И.О.")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("                                                                            Проверил:  Фамилия И.О.")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.left_indent = Cm(-0.63)
add = text.add_run("\n\n\n\n")
add.font.size = Pt(12)
add = text.add_run("\n\n\n\n\n")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = text.add_run("Тамбов 2001")
add.font.size = Pt(14)

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.25)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Цель работы:")
add.font.size = Pt(14)
add.bold = True

text = doc.add_paragraph()
p_fmt = text.paragraph_format
p_fmt.first_line_indent = Cm(1.25)
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
add = text.add_run("Зачем именно нужна данная работа. Что она позволяет освоить")
add.font.size = Pt(14)

doc.save("resources/Структура отчета.docx")