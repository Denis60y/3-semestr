import docx
import docx.enum.text
from docx.shared import Pt

def clear_cell_paragraphs(cell):
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'

table = doc.add_table(6, 2)
table.style = 'Table Grid'

for i in range(5):
    table.cell(i+1,0).merge(table.cell(i+1,1))

# def new_parafraph(cell, text, fontSize, bolt, italic, underline):
#     head = cell.add_paragraph()
#     p_fmt = head.paragraph_format
#     p_fmt.space_after = Pt(10)
#     add = head.add_run(text)
#     add.font.name = 'Times New Roman'
#     add.font.size = Pt(fontSize)
#     add.bold = bolt
#     add.font.italic = italic
#     add.underline = underline

#ячейка номер 1

cell = table.cell(0, 0)

clear_cell_paragraphs(cell)

head = cell.add_paragraph()
p_fmt = head.paragraph_format
p_fmt.space_after = Pt(10)
add = head.add_run("Дата поступления")
add.font.size = Pt(12)

void = cell.add_paragraph()
add = void.add_run("")
add.font.size = Pt(12)

dop = cell.add_paragraph()
add = dop.add_run("(заполняется Федеральной службой  по интеллектуальной собственности)")
add.font.size = Pt(9)
add.font.italic = True

#ячейка номер 2

cell = table.cell(0, 1)

clear_cell_paragraphs(cell)

head = cell.add_paragraph()
p_fmt = head.paragraph_format
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = head.add_run("В Федеральную службу\nпо интеллектуальной собственности")
add.font.size = Pt(9)
add.bold = True


dop = cell.add_paragraph()
p_fmt = dop.paragraph_format
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = dop.add_run("Бережковская наб., д. 30, корп. 1, \nг. Москва, Г-59, ГСП-3, 125993, \nРоссийская Федерация")
add.font.size = Pt(9)

#ячейка номер 3

cell = table.cell(1, 0)

clear_cell_paragraphs(cell)

head = cell.add_paragraph()
p_fmt = head.paragraph_format
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = head.add_run("Согласие на указание сведений об авторе в заявлении на государственную регистрацию программы для ЭВМ или базы данных")
add.font.size = Pt(13)
add.bold = True

#ячейка номер 4

cell = table.cell(2, 0)

clear_cell_paragraphs(cell)
head = cell.add_paragraph()
add = head.add_run("Заявка  №  ____________________________________________________________")
add.font.size = Pt(12)

head = cell.add_paragraph()
p_fmt = head.paragraph_format
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
p_fmt.space_after = Pt(10)
add = head.add_run("(указывается при наличии регистрационного номера заявки)")
add.font.size = Pt(7)
add.font.italic = True

dop = cell.add_paragraph()
p_fmt = head.paragraph_format
p_fmt.space_after = Pt(10)
add = dop.add_run("на государственную регистрацию: ")
add.font.size = Pt(12)
add = dop.add_run("☒ ")
add.font.size = Pt(10)
add = dop.add_run("Программы для ЭВМ ☐ Базы данных, государственная регистрация которой осуществляется в соответствии с пунктом 4 статьи 1259 Кодекса ☐ Базы данных, государственная регистрация которой осуществляется в соответствии с пунктом 3 статьи 1334 Кодекса ")
add.font.size = Pt(9)
add.bold = True
add = dop.add_run("(Отметить знаком «Х» вид результата интеллектуальной деятельности)")
add.font.size = Pt(9)
add.font.italic = True

dop = cell.add_paragraph()
add = dop.add_run("Название:")
add.font.size = Pt(12)
add = dop.add_run(" «Название программы ЭВМ»")
add.font.size = Pt(12)
add.underline = True

dop = cell.add_paragraph()
p_fmt = dop.paragraph_format
p_fmt.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
add = dop.add_run("(указывается в  соответствии с графой 1 заявления о государственной регистрации программы для ЭВМ или базы данных)")
add.font.size = Pt(8)
add.font.italic = True

#ячейка номер 5

cell = table.cell(3, 0)

clear_cell_paragraphs(cell)

head = cell.add_paragraph()
add = head.add_run("Правообладатель (и) (Заявитель)(и) ")
add.font.size = Pt(12)
add = head.add_run("(указываются фамилия, имя, отчество (последнее – при наличии), место жительства физического лица, наименование, место нахождения, основной государственный регистрационный номер (ОГРН)  и идентификационный номер налогоплательщика (ИНН) юридического лица) ")
add.font.size = Pt(9)
add.font.italic = True

head = cell.add_paragraph()
add = head.add_run("Физ. Лицо -  указываются фамилия, имя, отчество (последнее – при наличии), место жительства")
add.font.size = Pt(12)
add.underline = True

dop = cell.add_paragraph()
add = dop.add_run("Подтверждаю согласие на указание обо мне, как авторе, следующих сведений в графе 7А заявления на государственную регистрацию данной программы для ЭВМ или базы данных.")
add.font.size = Pt(11)
add.font.italic = True
add.bold = True

dop = cell.add_paragraph()
add = dop.add_run("7А. СВЕДЕНИЯ ОБ АВТОРЕ:")
add.font.size = Pt(10)
add.font.italic = True

dop = cell.add_paragraph()
add = dop.add_run("Фамилия имя отчество:")
add.font.size = Pt(10)
add.font.italic = True
add = dop.add_run("  Фамилия Имя Отчество")
add.font.size = Pt(12)
add.bold = True

dop = cell.add_paragraph()
add = dop.add_run("Фамилия имя отчество:")
add.font.size = Pt(10)
add.font.italic = True
add = dop.add_run("  11_   ")
add.font.size = Pt(10)
add.bold = True
add.underline = True
add = dop.add_run("месяц: ")
add.font.size = Pt(10)
add.font.italic = True
add = dop.add_run("  11   ")
add.font.size = Pt(10)
add.bold = True
add.underline = True
add = dop.add_run("год: ")
add.font.size = Pt(10)
add.font.italic = True
add = dop.add_run("  2001   ")
add.font.size = Pt(10)
add.bold = True
add.underline = True
add = dop.add_run("Гражданство:")
add.font.size = Pt(10)
add.font.italic = True
add = dop.add_run("_Гражданство_")
add.font.size = Pt(10)
add.bold = True
add.underline = True

dop = cell.add_paragraph()
add = dop.add_run("Автор согласен с обработкой указанных персональных данных, необходимой для исполнения полномочий федеральных органов исполнительной власти, участвующих в предоставлении государственных услуг, предусмотренных Федеральным законом от 27 июля 2010 года № 210-ФЗ «Об организации предоставления государственных и муниципальных услуг», включая регистрацию субъекта персональных данных на едином портале государственных и муниципальных услуг и (или) региональных порталах государственных и муниципальных услуг и в течение срока действия исключительного права на регистрируемый объект.")
add.font.size = Pt(10)
add.bold = True

dop = cell.add_paragraph()
add = dop.add_run("Место постоянного жительства, включая указание страны:")
add.font.size = Pt(10)
add.font.italic = True

dop = cell.add_paragraph()
add = dop.add_run("Место постоянного жительства")
add.font.size = Pt(12)
add.bold = True

dop = cell.add_paragraph()
add = dop.add_run("Краткое описание творческого вклада автора при создании регистрируемой программы для ЭВМ или базы данных:")
add.font.size = Pt(10)
add.font.italic = True

dop = cell.add_paragraph()
add = dop.add_run("Написание исходного текста программы – 0 %")
add.font.size = Pt(11)
add.bold = True

void = cell.add_paragraph()
add = void.add_run("")
add.font.size = Pt(10)

dop = cell.add_paragraph()
add = dop.add_run("При  публикации сведений о государственной регистрации программы для ЭВМ или базы данных автор просит: (отметить [X])")
add.font.size = Pt(10)
add.font.italic = True

#ячейка номер 6

cell = table.cell(4, 0)

clear_cell_paragraphs(cell)

head = cell.add_paragraph()
add = head.add_run("Подпись автора:                                                                     И.О. Фамилия")
add.font.size = Pt(12)

void = cell.add_paragraph()
add = void.add_run("")
add.font.size = Pt(10)

dop = cell.add_paragraph()
add = dop.add_run("(подпись должна быть расшифрована)")
add.font.size = Pt(9)
add.font.italic = True

#ячейка номер 7

cell = table.cell(5, 0)

clear_cell_paragraphs(cell)

head = cell.add_paragraph()
add = head.add_run("Подпись(и) правообладателя(ей) или его (их) представителя(ей)")
add.font.size = Pt(12)

void = cell.add_paragraph()
add = void.add_run("")
add.font.size = Pt(12)

dop = cell.add_paragraph()
add = dop.add_run("_________Должность____________________________И.О. Фамилия_______________")
add.font.size = Pt(11)
add.underline = True

dop = cell.add_paragraph()
p_fmt = dop.paragraph_format
p_fmt.space_after = Pt(10)
add = dop.add_run(" (от имени юридического лица заявление подписывается руководителем организации или иным лицом, уполномоченным на это в установленном законодательством Российской Федерации порядке, с указанием его должности, подпись удостоверяется печатью юридического лица при наличии печати. Подпись любого лица должна быть расшифрована с указанием фамилии и инициалов и даты подписания заявления)")
add.font.size = Pt(7)
add.font.italic = True

doc.save("resources/Согласие на указание сведений об авторе.docx")